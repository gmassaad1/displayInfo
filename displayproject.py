from email.utils import decode_rfc2231
import tkinter as tk
from PIL import Image, ImageTk
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
import re

class ImageShufflerGUI:
    def _init_(self, root):
        self.root = root
        self.images = []
        self.current_image = None
        self.words = []

        self.load_images()
        self.load_words()
        print(self.words)

        self.image_label = tk.Label(root)
        self.image_label.pack(side=tk.RIGHT, padx=0, fill=tk.BOTH, expand=True)

        self.text_textbox = tk.Text(root, wrap=tk.WORD, font=("Arial", 14))
        self.text_textbox.pack(side=tk.LEFT, padx=0, pady=0, fill=tk.BOTH, expand=True)

        # Remove border and cursor from the Text widget
        self.text_textbox.config(borderwidth=0, highlightthickness=0, relief=tk.FLAT, insertofftime=0)

        # Make the Text widget read-only
        self.text_textbox.config(state="disabled")

        self.current_image_index = 0
        self.root.bind("<Configure>", self.on_window_resize)
        self.root.update()  # Ensure the window is displayed before getting width and height
        self.update_display()

    def load_images(self):
        folder_path = "images"
        if not os.path.exists(folder_path):
            return

        for filename in os.listdir(folder_path):
            if filename.endswith(".jpg") or filename.endswith(".png") or filename.endswith(".jpeg"):
                image_path = os.path.join(folder_path, filename)
                self.images.append(image_path)

    def load_words(self):
        word_file = "words.docx"
        if not os.path.exists(word_file):
            return

        doc = Document(word_file)

        content = []
        for paragraph in doc.paragraphs:
            text = ""
            for run in paragraph.runs:
                text += run.text
                print(text)
                if run.text.endswith("\n"):
                    content.append(text.rstrip("\n"))
                    text = ""
            if text:
                content.append(text.rstrip("\n"))
        print(content)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                text = run.text
                self.words.append((text, run.bold, run.italic, run.underline, run.font.size, run.font.color))
            if text:
                self.words.append(("\n",None,None,None,None,None))

    def update_display(self):
        if len(self.images) == 0 or len(self.words) == 0:
            return
        if self.current_image_index >= len(self.images):
            self.current_image_index = 0

        image_path = self.images[self.current_image_index]
        self.current_image = Image.open(image_path)

        self.update_image_size()

        self.current_image = ImageTk.PhotoImage(self.current_image)
        self.image_label.config(image=self.current_image)

        self.display_words()

        self.current_image_index += 1

        self.root.after(3000, self.update_display)

    def display_words(self):
        if self.words:
            self.text_textbox.config(state="normal")
            self.text_textbox.delete("1.0", tk.END)

            for word in self.words:
                run_text, run_bold, run_italic, run_underline, run_font_size, run_font_color = word
                if run_text=="\n":
                    self.text_textbox.insert(tk.END, run_text)
                    start_index = tk.END + "-%dc" % len(run_text)
                    end_index = tk.END
                if run_text.strip():
                    self.text_textbox.insert(tk.END, run_text)

                    start_index = tk.END + "-%dc" % len(run_text)
                    end_index = tk.END

                    self.text_textbox.tag_add("bold", start_index, end_index)
                    if run_bold:
                        self.text_textbox.tag_configure("bold", font=("Arial",  12, "bold"))

                    self.text_textbox.tag_add("italic", start_index, end_index)
                    if run_italic:
                        self.text_textbox.tag_configure("italic", font=("Arial",  12, "italic"))

                    self.text_textbox.tag_add("underline", start_index, end_index)
                    if run_underline:
                        self.text_textbox.tag_configure("underline", font=("Arial", 12, "underline"))

                    self.text_textbox.tag_add("color", start_index, end_index)
                    if run_font_color:
                        self.text_textbox.tag_configure("color", foreground="#%02x%02x%02x" % (
                            run_font_color.rgb[0] if run_font_color.rgb else 0,
                            run_font_color.rgb[1] if run_font_color.rgb else 0,
                            run_font_color.rgb[2] if run_font_color.rgb else 0
                        ))

            self.text_textbox.config(state="disabled")

    def update_image_size(self, event=None):
        window_width = self.root.winfo_width()
        window_height = self.root.winfo_height()
        image_width = int(window_width / 2)
        image_height = window_height

        if self.current_image is not None:
            try:
                self.current_image = self.current_image.resize((image_width, image_height))
            except AttributeError:
                # Handle the AttributeError exception
                print("AttributeError: 'PhotoImage' object has no attribute 'resize'")
                # Additional error handling code or fallback behavior
    def on_window_resize(self, event):
        self.update_image_size()
        self.image_label.config(image=self.current_image)
        
def GUIStart():
    # Create the root window
    root = tk.Tk()
    root.title("Image Shuffler")
    root.attributes("-fullscreen", True)  # Set fullscreen attribute

    # Create the GUI
    app = ImageShufflerGUI(root)
    
    # Start the main event loop
    root.mainloop()
